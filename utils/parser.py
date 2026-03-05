# utils/parser.py
# Resume text extraction and structured information parsing

import re
import io
import os
import sys

# ── PDF / DOCX extraction ──────────────────────────────────────────────────────

def extract_text_from_pdf(filepath: str) -> str:
    """Extract raw text from a PDF file using PyPDF2."""
    try:
        import PyPDF2
        text = []
        with open(filepath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        return "\n".join(text)
    except Exception as e:
        print(f"[parser] PDF extraction error: {e}", file=sys.stderr)
        return ""


def extract_text_from_docx(filepath: str) -> str:
    """Extract raw text from a DOCX file using python-docx."""
    try:
        from docx import Document
        doc = Document(filepath)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract table content
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        print(f"[parser] DOCX extraction error: {e}", file=sys.stderr)
        return ""


def extract_text(filepath: str) -> str:
    """Route file to correct extractor based on extension."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(filepath)
    else:
        # Try reading as plain text fallback
        try:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception:
            return ""


# ── Contact Information Extraction ────────────────────────────────────────────

def extract_email(text: str) -> str:
    """Extract email address using regex."""
    pattern = r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}"
    matches = re.findall(pattern, text)
    return matches[0] if matches else ""


def extract_phone(text: str) -> str:
    """Extract phone number in various formats."""
    patterns = [
        r"(?:\+91[\s\-]?)?(?:\(\d{3}\)|\d{3})[\s\-]?\d{3}[\s\-]?\d{4}",
        r"\+?\d[\d\s\-().]{8,15}\d",
        r"\b\d{10}\b",
    ]
    for pattern in patterns:
        matches = re.findall(pattern, text)
        if matches:
            return matches[0].strip()
    return ""


def extract_name(text: str) -> str:
    """
    Heuristic name extraction:
    The name is usually the first non-empty line that is not an email/phone/URL
    and is relatively short (likely a person's name).
    """
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    skip_patterns = [
        r"@", r"http", r"\d{5,}", r"resume", r"curriculum", r"vitae",
        r"^(address|email|phone|mobile|tel|linkedin|github)"
    ]
    for line in lines[:8]:  # Search only in first 8 lines
        lower = line.lower()
        skip = any(re.search(p, lower) for p in skip_patterns)
        if not skip and 2 <= len(line.split()) <= 5 and len(line) < 60:
            # Likely a name (2-5 words, no special keywords)
            return line.title()
    return "Not Found"


def extract_linkedin(text: str) -> str:
    """Extract LinkedIn profile URL."""
    pattern = r"(?:linkedin\.com/in/|linkedin\.com/pub/)[\w\-]+"
    matches = re.findall(pattern, text, re.IGNORECASE)
    return "https://" + matches[0] if matches else ""


def extract_github(text: str) -> str:
    """Extract GitHub profile URL."""
    pattern = r"(?:github\.com/)[\w\-]+"
    matches = re.findall(pattern, text, re.IGNORECASE)
    return "https://" + matches[0] if matches else ""


# ── Section Segmentation ───────────────────────────────────────────────────────

SECTION_HEADERS = {
    "experience": [
        "experience", "work experience", "professional experience",
        "employment", "work history", "career", "internship"
    ],
    "education": [
        "education", "academic background", "qualifications",
        "academic qualifications", "educational background", "degrees"
    ],
    "skills": [
        "skills", "technical skills", "core competencies",
        "competencies", "technologies", "tools", "expertise"
    ],
    "projects": [
        "projects", "personal projects", "key projects",
        "academic projects", "portfolio"
    ],
    "certifications": [
        "certifications", "certificates", "courses",
        "training", "licenses"
    ],
    "summary": [
        "summary", "objective", "profile", "about me",
        "professional summary", "career objective", "overview"
    ],
    "achievements": [
        "achievements", "accomplishments", "awards",
        "honors", "recognition"
    ]
}


def identify_sections(text: str) -> dict:
    """
    Split resume text into named sections using header detection.
    Returns a dict mapping section_name -> section_text.
    """
    lines = text.split("\n")
    sections = {}
    current_section = "header"
    current_content = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            current_content.append("")
            continue

        detected = _detect_section_header(stripped)
        if detected:
            # Save the previous section
            sections[current_section] = "\n".join(current_content).strip()
            current_section = detected
            current_content = []
        else:
            current_content.append(stripped)

    # Save last section
    sections[current_section] = "\n".join(current_content).strip()
    return sections


def _detect_section_header(line: str) -> str | None:
    """Return section key if the line looks like a section header."""
    clean = line.lower().strip().rstrip(":").strip()
    for section, aliases in SECTION_HEADERS.items():
        if clean in aliases:
            return section
        # Also match if line starts with or is mostly the alias
        for alias in aliases:
            if clean.startswith(alias) and len(clean) < len(alias) + 12:
                return section
    return None


# ── Education Parsing ──────────────────────────────────────────────────────────

DEGREE_KEYWORDS = [
    "b.e", "b.tech", "be", "btech", "bachelor", "b.s", "bs",
    "m.tech", "mtech", "master", "m.s", "ms", "mba", "phd",
    "ph.d", "diploma", "associate", "b.sc", "m.sc", "bsc", "msc",
    "b.com", "m.com", "bca", "mca", "b.ca", "m.ca"
]

YEAR_PATTERN = re.compile(r"\b(19|20)\d{2}\b")


def extract_education(text: str, sections: dict) -> list:
    """Extract education entries from education section or full text."""
    source = sections.get("education", text)
    entries = []
    lines = [l.strip() for l in source.split("\n") if l.strip()]

    for line in lines:
        lower = line.lower()
        has_degree = any(kw in lower for kw in DEGREE_KEYWORDS)
        has_year = bool(YEAR_PATTERN.search(line))
        if has_degree or has_year:
            entries.append(line)

    return entries[:5]  # Return top 5 education entries


# ── Experience Parsing ─────────────────────────────────────────────────────────

def extract_experience(sections: dict) -> list:
    """Extract work experience bullet points."""
    source = sections.get("experience", "")
    if not source:
        return []

    lines = [l.strip() for l in source.split("\n") if l.strip()]
    entries = []
    for line in lines:
        # Keep lines that look like job titles, companies, or bullet points
        if len(line) > 10:
            entries.append(line)

    return entries[:20]  # Top 20 lines


# ── Skills Extraction ──────────────────────────────────────────────────────────

def extract_skills(text: str, skills_db: dict) -> dict:
    """
    Match resume text against the skills database.
    Returns categorized dict of found skills.
    """
    text_lower = text.lower()
    found = {}

    for category, skill_list in skills_db.items():
        matched = []
        for skill in skill_list:
            # Use word-boundary matching to avoid partial matches
            pattern = r"\b" + re.escape(skill) + r"\b"
            if re.search(pattern, text_lower):
                matched.append(skill)
        if matched:
            found[category] = matched

    return found


def get_flat_skills(skills_dict: dict) -> list:
    """Flatten categorized skills dict to a simple list."""
    flat = []
    for skills in skills_dict.values():
        flat.extend(skills)
    return flat


# ── Keyword Extraction ─────────────────────────────────────────────────────────

def extract_keywords(text: str, top_n: int = 30) -> list:
    """
    Simple TF-IDF-style keyword extraction without spaCy dependency.
    Uses stop words and frequency analysis.
    """
    STOP_WORDS = {
        "the", "a", "an", "and", "or", "but", "in", "on", "at", "to",
        "for", "of", "with", "by", "from", "is", "are", "was", "were",
        "be", "been", "being", "have", "has", "had", "do", "does", "did",
        "will", "would", "could", "should", "may", "might", "shall",
        "this", "that", "these", "those", "i", "my", "me", "we", "our",
        "you", "your", "he", "she", "it", "they", "their", "its",
        "as", "if", "then", "than", "so", "up", "out", "about",
        "into", "through", "during", "including", "until", "while",
        "not", "no", "nor", "can", "also", "both", "each", "other",
        "more", "most", "very", "just", "here", "there", "when", "where",
        "who", "which", "how", "all", "any", "some", "such", "new",
        "per", "etc", "using", "used", "use", "within"
    }

    words = re.findall(r"\b[a-zA-Z][a-zA-Z+#./\-]{1,30}\b", text.lower())
    freq = {}
    for word in words:
        if word not in STOP_WORDS and len(word) > 2:
            freq[word] = freq.get(word, 0) + 1

    sorted_keywords = sorted(freq.items(), key=lambda x: x[1], reverse=True)
    return [kw for kw, _ in sorted_keywords[:top_n]]


# ── Projects & Certifications ──────────────────────────────────────────────────

def extract_projects(sections: dict) -> list:
    """Extract project titles from the projects section."""
    source = sections.get("projects", "")
    if not source:
        return []
    lines = [l.strip() for l in source.split("\n") if l.strip() and len(l.strip()) > 5]
    return lines[:10]


def extract_certifications(sections: dict) -> list:
    """Extract certification names."""
    source = sections.get("certifications", "")
    if not source:
        return []
    lines = [l.strip() for l in source.split("\n") if l.strip() and len(l.strip()) > 5]
    return lines[:10]


# ── Master Parse Function ──────────────────────────────────────────────────────

def parse_resume(filepath: str, skills_db: dict) -> dict:
    """
    Full pipeline: extract text → parse all sections → return structured dict.
    """
    raw_text = extract_text(filepath)
    if not raw_text:
        return {"error": "Could not extract text from resume. Please check the file."}

    sections = identify_sections(raw_text)

    # Combine all text for skill matching
    full_text = raw_text

    parsed = {
        "raw_text": raw_text,
        "name": extract_name(raw_text),
        "email": extract_email(raw_text),
        "phone": extract_phone(raw_text),
        "linkedin": extract_linkedin(raw_text),
        "github": extract_github(raw_text),
        "skills": extract_skills(full_text, skills_db),
        "flat_skills": get_flat_skills(extract_skills(full_text, skills_db)),
        "education": extract_education(raw_text, sections),
        "experience": extract_experience(sections),
        "projects": extract_projects(sections),
        "certifications": extract_certifications(sections),
        "keywords": extract_keywords(raw_text),
        "sections_found": [k for k, v in sections.items() if v and k != "header"],
        "sections": sections,
        "word_count": len(raw_text.split()),
        "char_count": len(raw_text)
    }

    return parsed
